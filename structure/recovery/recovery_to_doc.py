import os
from copy import deepcopy
import shutil
from docx import Document
from docx import shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

from structure.recovery.table_process import HtmlToDocx

from utils.logging import get_logger
logger = get_logger()


def convert_info_docx(img, res, save_folder, img_name):
    figure_dir = 'static/output/figure/'
    for filename in os.listdir(figure_dir):
            file_path = os.path.join(figure_dir, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print('Failed to delete %s. Reason: %s' % (file_path, e))
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = shared.Pt(6.5)
    html_file = open('static/output/html/docx_output.html', 'w')
            # text_file.write(text)

    flag = 1
    for i, region in enumerate(res):
        img_idx = region['img_idx']
        if flag == 2 and region['layout'] == 'single':
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '1')
            flag = 1
        elif flag == 1 and region['layout'] == 'double':
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
            flag = 2

        if region['type'].lower() == 'figure':
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(excel_save_folder,
                                    '{}_{}.jpg'.format(region['bbox'], img_idx))
            shutil.copy2(img_path, figure_dir)
            
            html_img_path = '../figure/' + img_path.split('/')[-1]
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph_pic.add_run("")
            if flag == 1:
                run.add_picture(img_path, width=shared.Inches(5))
            elif flag == 2:
                run.add_picture(img_path, width=shared.Inches(2))
            html_line = '<img src="{}" >'.format(html_img_path)
            html_file.write(html_line)
        elif region['type'].lower() == 'title':
            doc.add_heading(region['res'][0]['text'])
            html_line = '<h2>{}</h2>'.format(region['res'][0]['text'])
            html_file.write(html_line)
        elif region['type'].lower() == 'table':
            parser = HtmlToDocx()
            parser.table_style = 'TableGrid'
            parser.handle_table(region['res']['html'], doc)
            html_file.write(region['res']['html'])
        else:
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            text_line = ''
            for i, line in enumerate(region['res']):
                if i == 0:
                    paragraph_format.first_line_indent = shared.Inches(0.25)
                text_run = paragraph.add_run(line['text'] + ' ')
                text_line = text_line + line['text'] + ' '
                text_run.font.size = shared.Pt(10)
                if len(line['text']) > 5:
                    html_line = '<p>{}</p>'.format(text_line)
                    html_file.write(html_line)
                    text_line = ''
    table_style = "<style> \
                table { \
                width:100%; \
                border:1px solid black; \
                } \
                th, td { \
                border:1px solid black; \
                } \
                </style>"
    html_file.write(table_style) 
    html_file.close()
    # save to docx
    docx_save_path = os.path.join(save_folder, img_name)
    docx_path = os.path.join(docx_save_path, '{}_ocr.docx'.format(img_name))
    doc.save(docx_path)
    logger.info('docx save to {}'.format(docx_path))
    return docx_path


def sorted_layout_boxes(res, w):
    """
    Sort text boxes in order from top to bottom, left to right
    args:
        res(list):structure results
    return:
        sorted results(list)
    """
    num_boxes = len(res)
    if num_boxes == 1:
        res[0]['layout'] = 'single'
        return res

    sorted_boxes = sorted(res, key=lambda x: (x['bbox'][1], x['bbox'][0]))
    _boxes = list(sorted_boxes)

    new_res = []
    res_left = []
    res_right = []
    i = 0

    while True:
        if i >= num_boxes:
            break
        if i == num_boxes - 1:
            if _boxes[i]['bbox'][1] > _boxes[i - 1]['bbox'][3] and _boxes[i][
                    'bbox'][0] < w / 2 and _boxes[i]['bbox'][2] > w / 2:
                new_res += res_left
                new_res += res_right
                _boxes[i]['layout'] = 'single'
                new_res.append(_boxes[i])
            else:
                if _boxes[i]['bbox'][2] > w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_right.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
                elif _boxes[i]['bbox'][0] < w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_left.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
            res_left = []
            res_right = []
            break
        elif _boxes[i]['bbox'][0] < w / 4 and _boxes[i]['bbox'][2] < 3 * w / 4:
            _boxes[i]['layout'] = 'double'
            res_left.append(_boxes[i])
            i += 1
        elif _boxes[i]['bbox'][0] > w / 4 and _boxes[i]['bbox'][2] > w / 2:
            _boxes[i]['layout'] = 'double'
            res_right.append(_boxes[i])
            i += 1
        else:
            new_res += res_left
            new_res += res_right
            _boxes[i]['layout'] = 'single'
            new_res.append(_boxes[i])
            res_left = []
            res_right = []
            i += 1
    if res_left:
        new_res += res_left
    if res_right:
        new_res += res_right
    return new_res
